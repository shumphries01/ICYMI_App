# -*- coding: utf-8 -*-
"""
Created on Sat Jul 25 18:24:34 2020

@author: Sam Humphries

Algorithm that scrapes all relevant news articles and compiles on Word doc.
"""

import pandas as pd
import requests
from bs4 import BeautifulSoup

from newspaper import Article

import docx
from docx.enum.dml import MSO_THEME_COLOR_INDEX
from docx import Document
from docx.shared import Inches
from docx.shared import Pt


import re
import tldextract

import dateutil.parser as dp
import dateutil.tz as dtz
import datetime

######### Helper Functions ###############

##Define Time Checker Function
def timeChecker(date):
    date_string = date
    adatetime=dp.parse(date_string)
    adatetime = adatetime.replace(tzinfo=dtz.tzlocal())
    now=datetime.datetime.now(dtz.tzlocal())
    TwoDaysAgo=now-datetime.timedelta(days=2)
    if adatetime >= TwoDaysAgo:
        return True
    
##Define FindAuthor Function
def findAuthor(link):
    article = Article(link)
    article.download()
    article.parse()
    #return article.authors
    if len(article.authors) != 0:
        return (article.authors[0])
    else:
        return 'Anonymous'

##Defne Word Checker Function
def findWholeWord(w):
    return re.compile(r'\b({0})\b'.format(w), flags=re.IGNORECASE).search

##Define Hyperlink Function


def add_hyperlink(paragraph, text, url):
    # This gets access to the document.xml.rels file and gets a new relation id value
    part = paragraph.part
    r_id = part.relate_to(url, docx.opc.constants.RELATIONSHIP_TYPE.HYPERLINK, is_external=True)

    # Create the w:hyperlink tag and add needed values
    hyperlink = docx.oxml.shared.OxmlElement('w:hyperlink')
    hyperlink.set(docx.oxml.shared.qn('r:id'), r_id, )

    # Create a w:r element and a new w:rPr element
    new_run = docx.oxml.shared.OxmlElement('w:r')
    rPr = docx.oxml.shared.OxmlElement('w:rPr')

    # Join all the xml elements together add add the required text to the w:r element
    new_run.append(rPr)
    new_run.text = text
    hyperlink.append(new_run)

    # Create a new Run object and add the hyperlink into it
    r = paragraph.add_run ()
    r._r.append (hyperlink)

    # A workaround for the lack of a hyperlink style (doesn't go purple after using the link)
    # Delete this if using a template that has the hyperlink style in it
    r.font.color.theme_color = MSO_THEME_COLOR_INDEX.HYPERLINK
    r.font.underline = True
    r.font.bold = True

    return hyperlink



# =============================================================================
# if findWholeWord('climate')(description) is not None:
#     print('yea')
#     
# news_items[0]['description']

# =============================================================================
# word_set = set(['economy','Economic'])
# word_set
# phrase_set = set('chinas economic collapse'.split())
# if word_set.intersection(phrase_set):
#     print('go')
#     
# search_list = ['one', 'two', 'there']
# long_string = 'some one long two phrase three'
# if re.compile('|'.join(search_list),re.IGNORECASE).search(long_string):
#     print('go')
# =============================================================================
# =============================================================================
# =============================================================================
# =============================================================================
# article = Article(news_items[0]['link'])
# article.download()
# article.parse()
# article.keywords
# =============================================================================
# =============================================================================

# =============================================================================
# l = (news_items[0]['link'])
# print(l)
# 
# info = tldextract.extract(l)
# print(info.domain)
# =============================================================================
########### Start Algorithm ############

def scraper(url1,url2,url3,url4,url5, topic1 = False, topic2= False, topic3= False, topic4= False, topic5= False, topic6= False,
            topic7= False, topic8= False):
##Make List to Append Data
    news_items = []
    
    ##Set Category Lists
    
    AmComp_list = ['economy', 'economic','market']
    Arctic_list = ['arctic','polar']
    Asymmetric_list = ['attack', 'terrorist','terror','bomb','kill','qaeda','taliban','isis']
    Climate_list = ['climate','temperature','heat','floods']
    Energy_list = ['oil', 'energy','fuel','gas']
    NatSec_list = ['china','turmoil','Iran','coup']
    Nuclear_list = ['nuclear']
    USRussia_list = ['russia', 'putin']

    ##Enter URl and Perform Soup Magic
    #url = "http://feeds.bbci.co.uk/news/rss.xml" "https://www.theguardian.com/world/rss" ##BBC News
    #url_list = ["https://www.cnbc.com/id/100727362/device/rss/rss.html","http://www.npr.org/rss/rss.php?id=1004","http://feeds.washingtonpost.com/rss/world","https://aljazeera.com/xml/rss/all.xml","https://feeds.a.dj.com/rss/RSSWorldNews.xml","http://feeds.bbci.co.uk/news/rss.xml","https://rss.nytimes.com/services/xml/rss/nyt/World.xml"] ##NYT World
    #url_list = ["https://www.cnbc.com/id/100727362/device/rss/rss.html","http://www.npr.org/rss/rss.php?id=1004","https://aljazeera.com/xml/rss/all.xml","http://feeds.bbci.co.uk/news/rss.xml","https://rss.nytimes.com/services/xml/rss/nyt/World.xml"] ##NYT World
    url_list = [url1,url2,url3,url4,url5]
    #url_list = ["https://www.theguardian.com/world/rss","https://aljazeera.com/xml/rss/all.xml","https://feeds.a.dj.com/rss/RSSWorldNews.xml","http://feeds.bbci.co.uk/news/rss.xml","https://rss.nytimes.com/services/xml/rss/nyt/World.xml"]
    ##Main Loop
    
    for url in url_list:
        if (url != ""):
            resp = requests.get(url)
        
            soup = BeautifulSoup(resp.content, features = "xml")
        
            ##Inspect Unicode
            #print(soup.prettify())
        
            ##Inspect All Content
            items = soup.findAll('item')
            #print(items)
        
        
        
        ##Scrape HTML tags
            for item in items:
                date_string = item.pubDate.text
                if timeChecker(date_string) is True:
                    news_item = {}
                    news_item['title'] = item.title.text
                    news_item['description'] = item.description.text
                    news_item['link'] = item.link.text
                    info = tldextract.extract(item.link.text)
                    news_item['source'] = info.domain
                    news_item['pubDate'] = item.pubDate.text
                    news_item['author'] = findAuthor(item.link.text)
                    if re.compile('|'.join(AmComp_list),re.IGNORECASE).search(item.title.text):
                        news_item['topic'] = 'American Competitiveness'
                    elif re.compile('|'.join(Arctic_list),re.IGNORECASE).search(item.title.text):
                        news_item['topic'] = 'Arctic'
                    elif re.compile('|'.join(Asymmetric_list),re.IGNORECASE).search(item.title.text):
                        news_item['topic'] = 'Asymmetric Operations'
                    elif re.compile('|'.join(Climate_list),re.IGNORECASE).search(item.title.text):
                        news_item['topic'] = 'Climate Security'
                    elif re.compile('|'.join(Energy_list),re.IGNORECASE).search(item.title.text):
                        news_item['topic'] = 'Energy Security'
                    elif re.compile('|'.join(NatSec_list),re.IGNORECASE).search(item.title.text):
                        news_item['topic'] = 'National Security and Strategy'
                    elif re.compile('|'.join(Nuclear_list),re.IGNORECASE).search(item.title.text):
                        news_item['topic'] = 'Nuclear Security'
                    elif re.compile('|'.join(USRussia_list),re.IGNORECASE).search(item.title.text):
                        news_item['topic'] = 'US-Russia Relations'
                    else:
                        news_item['topic'] = 'Unknown'
                    news_items.append(news_item)
         
        # =============================================================================

    # =============================================================================
    ##Make DF
    df = pd.DataFrame(news_items, columns = ['title','description','link','source','pubDate','author','topic'])
    
    #Formalize Source Names
    df.replace(['npr','nytimes','aljazeera','cnbc','bbc'],
               ['NPR', 'The New York Times', 'Al Jazeera', 'CNBC','BBC'],inplace = True)
            
            
    
    
    ##################
    ##Write to Word Doc
    ##################
    
    ##Inititate Document
    document = Document()
    
    document.add_heading('ASP: In Case You Missed It...', 0)
    
    
    ##Sort
    new_df = df.sort_values('topic')
    ##Drop All Unknowns
    indexNames = new_df[new_df['topic'] == 'Unknown'].index
    new_df.drop(indexNames , inplace=True)
    ##Filter Based on Topic Selections
    if (topic1==False):
        new_df.drop(new_df[new_df['topic'] == 'American Competitiveness'].index , inplace=True)
    if (topic2==False):
        new_df.drop(new_df[new_df['topic'] == 'Arctic'].index , inplace=True)
    if (topic3==False):
        new_df.drop(new_df[new_df['topic'] == 'Asymmetric Operations'].index , inplace=True)
    if (topic4==False):
        new_df.drop(new_df[new_df['topic'] == 'Climate Security'].index , inplace=True)
    if (topic5==False):
        new_df.drop(new_df[new_df['topic'] == 'Energy Security'].index , inplace=True)
    if (topic6==False):
        new_df.drop(new_df[new_df['topic'] == 'National Security and Strategy'].index , inplace=True)
    if (topic7==False):
        new_df.drop(new_df[new_df['topic'] == 'Nuclear Security'].index , inplace=True)   
    if (topic8==False):
        new_df.drop(new_df[new_df['topic'] == 'US-Russia Relations'].index , inplace=True)
    ##Reset Index
    new_df = new_df.reset_index(drop=True)
    
    
    ##Main Loop for Word Doc Builder
    
    row_iterator = new_df.iterrows()
    _, last = next(row_iterator)
    
    for index, row in new_df.iterrows():
        if index == 0:
            header = document.add_heading(row['topic'], 1)
            header.paragraph_format.space_after = Pt(26)
            t = document.add_paragraph()
            add_hyperlink(t, row['title'],row['link'])
            t.paragraph_format.space_after = Pt(0.5)
            a = document.add_paragraph()
            a.add_run(row['author']).italic = True
            a.add_run(' | ' + row['source']).italic = True
            a.paragraph_format.space_after = Pt(0.5)
            d = document.add_paragraph(row['description'])
            #d.paragraph_format.space_after = Pt(45)
        elif row['topic']!=last['topic']:
            header = document.add_heading(row['topic'], 1)
            header.paragraph_format.space_before = Pt(55)
            header.paragraph_format.space_after = Pt(26)
            t = document.add_paragraph()
            add_hyperlink(t, row['title'],row['link'])
            t.paragraph_format.space_after = Pt(0.5)
            a = document.add_paragraph()
            a.add_run(row['author']).italic = True
            a.add_run(' | ' + row['source']).italic = True
            a.paragraph_format.space_after = Pt(0.5)
            document.add_paragraph(row['description'])
        else:
            t = document.add_paragraph()
            add_hyperlink(t, row['title'],row['link'])
            t.paragraph_format.space_before = Pt(5)
            t.paragraph_format.space_after = Pt(0.5)
            a = document.add_paragraph()
            a.add_run(row['author']).italic = True
            a.add_run(' | ' + row['source']).italic = True
            a.paragraph_format.space_after = Pt(0.5)
            document.add_paragraph(row['description'])
        last = row
                  
        
    
    # =============================================================================
    #     if row['topic'] is not row-1['topic']:
    #         document.add_heading(row['topic'], 1)
    #         document.add_paragraph(item.title.text, style='Intense Quote')
    #         document.add_paragraph(news_item['author']).italic=True
    #     else:
    #         document.add_paragraph(item.title.text, style='Intense Quote')
    #         document.add_paragraph(news_item['author']).italic=True
    # =============================================================================
    
    #df.head()
    
    print ('loading...')
    
    #new_df.to_csv('demo27AUG.csv',index = False, encoding = 'utf-8')
    
    document.save('demoPyPI.docx')
    
#scraper()
